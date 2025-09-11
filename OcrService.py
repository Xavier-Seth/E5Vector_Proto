import os
import re
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from docx import Document as DocxDocument
import pandas as pd
from pathlib import Path
from typing import List, Optional


class OcrService:
    """
    Lightweight OCR/extractor for PDFs, images, DOCX, XLS/XLSX, and TXT.

    ⚡ Performance defaults:
      - PDF rasterization at 200 DPI
      - Only the first 2 pages are OCR’d (configurable)
      - Tesseract OEM/PSM tuned for general documents

    You can tweak behavior via constructor args:
      OcrService(
        tesseract_cmd='tesseract',
        poppler_path=None,
        lang='eng',           # e.g. 'eng+fil'
        oem=3,                # LSTM
        psm=6,                # uniform block
        pdf_dpi=200,          # 150–200 is usually enough for categorization
        pdf_max_pages=2       # OCR only first N pages to speed up
      )
    """

    def __init__(
        self,
        tesseract_cmd: str = 'tesseract',
        poppler_path: Optional[str] = None,   # set if poppler isn't on PATH (Windows)
        lang: str = 'eng',                    # use 'eng+fil' if you want Filipino too
        oem: int = 3,                         # default LSTM
        psm: int = 6,                         # assume a uniform block of text
        pdf_dpi: int = 200,                   # ↓ from 300 to 200 for speed
        pdf_max_pages: int = 2            # only first 1–2 pages for classification
    ):
        self.tesseract_cmd = tesseract_cmd
        pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd

        self.poppler_path = poppler_path
        self.lang = lang
        self.tess_config = f'--oem {oem} --psm {psm}'

        # speed knobs
        self.pdf_dpi = int(pdf_dpi)
        self.pdf_max_pages = max(1, int(pdf_max_pages))

    # ── Public API ────────────────────────────────────────────────────────────
    def extract_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext in {'.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp'}:
            return self._extract_from_image(file_path)
        elif ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif ext == '.docx':
            return self._extract_from_docx(file_path)
        elif ext in {'.xls', '.xlsx'}:
            return self._extract_from_excel(file_path)
        elif ext == '.txt':
            return self._extract_from_txt(file_path)
        return ''

    # Optional helper: split long text into smaller snippets (3–10 lines each)
    def split_to_snippets(self, text: str, max_words: int = 80) -> List[str]:
        text = self.clean_text(text)
        # try to split on section headers like I., II., 1., 2., etc.
        parts = re.split(r'(?=\s(?:[IVXLC]+\.|[0-9]+\.)\s)', text)
        chunks: List[str] = []
        if len(parts) >= 3:
            for p in parts:
                p = p.strip()
                if p:
                    chunks.extend(self._split_by_words(p, max_words))
        else:
            chunks = self._split_by_words(text, max_words)
        # keep only meaningful snippets
        return [c for c in chunks if len(c.split()) >= 15]

    # ── Internal extractors ───────────────────────────────────────────────────
    def _extract_from_image(self, image_path: str) -> str:
        try:
            img = Image.open(image_path)
            text = pytesseract.image_to_string(img, lang=self.lang, config=self.tess_config)
            return self.clean_text(text)
        except Exception as e:
            print(f"Error reading image: {e}")
            return ''

    def _extract_from_pdf(self, pdf_path: str) -> str:
        text_parts = []
        try:
            # ⚡ Only render the first N pages at 200 DPI
            images = convert_from_path(
                pdf_path,
                dpi=self.pdf_dpi,
                first_page=1,
                last_page=self.pdf_max_pages,
                poppler_path=self.poppler_path  # set on Windows if needed
            )
            for img in images:
                page_txt = pytesseract.image_to_string(img, lang=self.lang, config=self.tess_config)
                text_parts.append(page_txt)
        except Exception as e:
            print(f"Error reading PDF: {e}")
        return self.clean_text("\n".join(text_parts))

    def _extract_from_docx(self, docx_path: str) -> str:
        try:
            doc = DocxDocument(docx_path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            table_rows = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_rows.append(row_text)
            text = "\n".join(paras + table_rows)
            return self.clean_text(text)
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ''

    def _extract_from_excel(self, excel_path: str) -> str:
        try:
            # engine will auto-select; dtype=str to avoid NaN
            dfs = pd.read_excel(excel_path, sheet_name=None, dtype=str)
            parts = []
            for _, df in dfs.items():
                df = df.fillna('')
                # Keep structure with newlines to help the embedding model
                lines = []
                for row in df.itertuples(index=False, name=None):
                    cells = [c for c in row if c]
                    if cells:
                        lines.append(" | ".join(cells))
                if lines:
                    parts.append("\n".join(lines))
            return self.clean_text("\n\n".join(parts))
        except Exception as e:
            print(f"Error reading Excel: {e}")
            return ''

    def _extract_from_txt(self, txt_path: str) -> str:
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                return self.clean_text(f.read())
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return ''

    # ── Utilities ────────────────────────────────────────────────────────────
    def clean_text(self, s: str) -> str:
        s = (s or '')
        s = s.replace('\x0c', ' ')      # page breaks
        s = s.replace('nan', ' ')
        s = re.sub(r'[ \t]+', ' ', s)   # collapse spaces
        s = re.sub(r'\s+\n', '\n', s)
        s = re.sub(r'\n{3,}', '\n\n', s)
        return s.strip()

    def _split_by_words(self, text: str, max_words: int) -> List[str]:
        words = text.split()
        chunks = []
        for i in range(0, len(words), max_words):
            chunk = " ".join(words[i:i+max_words]).strip()
            if chunk:
                chunks.append(chunk)
        return chunks
